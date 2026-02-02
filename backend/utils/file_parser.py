import os
import logging
from pathlib import Path
import zipfile
import statistics
import fitz  # PyMuPDF
from docx import Document

logger = logging.getLogger(__name__)


# ------------------------------------------------------------
# DOCX ORIGIN DETECTION (kept; uses python-docx + zipfile only)
# ------------------------------------------------------------

def detect_docx_origin(docx_path: str) -> str:
    """
    Heuristic classifier:
      - PDF_CONVERTED_DOCX
      - ORIGINAL_DOCX
      - DOC_TO_DOCX
    """
    doc = Document(docx_path)

    paras = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
    total_paras = len(paras)

    if total_paras == 0:
        return "ORIGINAL_DOCX"

    word_counts = [len(p.split()) for p in paras]
    short_lines = sum(1 for wc in word_counts if wc <= 3)
    short_ratio = short_lines / max(total_paras, 1)

    avg_words = sum(word_counts) / max(len(word_counts), 1)

    no_punct_lines = sum(1 for p in paras if p and p[-1] not in ".!?:;)")
    no_punct_ratio = no_punct_lines / max(total_paras, 1)

    table_count = len(doc.tables)

    xml_score = 0
    try:
        with zipfile.ZipFile(docx_path, "r") as z:
            doc_xml = z.read("word/document.xml").decode("utf-8", errors="ignore")

        if ("lastRenderedPageBreak" in doc_xml) or ('w:type="page"' in doc_xml):
            xml_score += 2

        if ("txbxContent" in doc_xml) or ("framePr" in doc_xml) or ("posH" in doc_xml) or ("posV" in doc_xml):
            xml_score += 2

        if doc_xml.count("<w:tbl") >= 3:
            xml_score += 1

        if doc_xml.count("<w:br") >= 20:
            xml_score += 1

    except Exception:
        pass

    score = xml_score

    if short_ratio >= 0.55:
        score += 2
    elif short_ratio >= 0.40:
        score += 1

    if avg_words < 6 and total_paras > 60:
        score += 1

    if no_punct_ratio >= 0.70 and total_paras > 50:
        score += 1

    if table_count >= 3:
        score += 1

    joined = "\n".join(paras).lower()
    if any(x in joined for x in ["adobe", "acrobat", "converted from pdf", "evaluation warning"]):
        score += 1

    if score >= 4:
        return "PDF_CONVERTED_DOCX"

    times_new_roman_hits = 0
    for p in doc.paragraphs:
        for r in p.runs:
            if r.font.name and "Times New Roman" in str(r.font.name):
                times_new_roman_hits += 1
                if times_new_roman_hits >= 10:
                    return "DOC_TO_DOCX"

    return "ORIGINAL_DOCX"


# ------------------------------------------------------------
# DOCX TEXT EXTRACTION (python-docx only)
# ------------------------------------------------------------

def read_docx_text(docx_path: str) -> str:
    """
    Safe DOCX extraction using python-docx.
    Includes paragraphs + tables, de-duped preserving order.
    """
    doc = Document(docx_path)

    lines = []

    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            lines.append(t)

    for tbl in doc.tables:
        for row in tbl.rows:
            # join row cells in reading order so it doesn't become "cell cell cell" noise
            row_text = []
            for cell in row.cells:
                ct = (cell.text or "").strip()
                if ct:
                    row_text.append(ct)
            if row_text:
                lines.append(" | ".join(row_text))

    # De-dupe preserve order
    seen = set()
    out = []
    for ln in lines:
        if ln not in seen:
            seen.add(ln)
            out.append(ln)

    return "\n".join(out).strip()


def _looks_bad_text(text: str) -> bool:
    """
    Small quality check:
      - too short
      - too many tiny lines (PDF-like wrapping)
    """
    if not text or len(text.strip()) < 200:
        return True

    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    if len(lines) < 5:
        return True

    short = sum(1 for ln in lines if len(ln.split()) <= 2)
    if short / max(len(lines), 1) > 0.55:
        return True

    return False


# ------------------------------------------------------------
# PDF LAYOUT-AWARE EXTRACTION (PyMuPDF only)
# ------------------------------------------------------------

def _page_two_col_split_x(page: fitz.Page, gap_quantile: float = 0.90, min_gap_ratio: float = 0.08):
    """
    Find a likely split-x using text block x-positions.

    Returns split_x or None if not confident.
    """
    blocks = page.get_text("blocks") or []
    # blocks: (x0, y0, x1, y1, text, block_no, block_type)
    x0s = sorted(b[0] for b in blocks if len(b) >= 5 and (b[4] or "").strip())
    if len(x0s) < 12:
        return None

    gaps = [b - a for a, b in zip(x0s, x0s[1:]) if (b - a) > 0]
    if not gaps:
        return None

    # robust gap threshold
    gaps_sorted = sorted(gaps)
    idx = int(max(0, min(len(gaps_sorted) - 1, round(gap_quantile * (len(gaps_sorted) - 1)))))
    largest_gap = gaps_sorted[idx]

    page_width = float(page.rect.width)
    if largest_gap < page_width * float(min_gap_ratio):
        return None

    # choose the midpoint of the *largest actual gap* (not just quantile)
    best_gap = 0.0
    split_x = None
    for a, b in zip(x0s, x0s[1:]):
        gap = b - a
        if gap > best_gap:
            best_gap = gap
            split_x = (a + b) / 2.0

    return split_x


def is_two_column_pymupdf(page: fitz.Page, min_gap_ratio: float = 0.08) -> bool:
    split_x = _page_two_col_split_x(page, min_gap_ratio=min_gap_ratio)
    if split_x is None:
        return False

    blocks = page.get_text("blocks") or []
    left = 0
    right = 0
    for b in blocks:
        if len(b) < 5:
            continue
        text = (b[4] or "").strip()
        if not text:
            continue
        x0, x1 = float(b[0]), float(b[2])
        x_center = (x0 + x1) / 2.0
        if x_center < split_x:
            left += 1
        else:
            right += 1

    return left > 0 and right > 0


def _sort_blocks_reading_order(blocks):
    """
    Sort blocks in natural reading order: by y0 then x0.
    """
    return sorted(blocks, key=lambda b: (round(float(b[1]), 1), round(float(b[0]), 1)))


def read_pdf_sequential(doc: fitz.Document) -> str:
    out = []
    for page in doc:
        blocks = page.get_text("blocks") or []
        blocks = [b for b in blocks if len(b) >= 5 and (b[4] or "").strip()]
        blocks = _sort_blocks_reading_order(blocks)
        out.append("\n".join((b[4] or "").strip() for b in blocks if (b[4] or "").strip()))
    return "\n\n".join(t for t in out if t.strip()).strip()


def read_pdf_two_column_left_then_right(doc: fitz.Document, header_top: float = 90.0, min_gap_ratio: float = 0.08) -> str:
    """
    Two-column layout handling:
      - header region (top area) read first in normal order
      - then left column top->bottom
      - then right column top->bottom
    """
    out_pages = []

    for page in doc:
        blocks = page.get_text("blocks") or []
        blocks = [b for b in blocks if len(b) >= 5 and (b[4] or "").strip()]
        if not blocks:
            continue

        split_x = _page_two_col_split_x(page, min_gap_ratio=min_gap_ratio)
        if split_x is None:
            # fallback
            blocks = _sort_blocks_reading_order(blocks)
            out_pages.append("\n".join((b[4] or "").strip() for b in blocks))
            continue

        header = []
        body = []
        for b in blocks:
            y0 = float(b[1])
            if y0 < header_top:
                header.append(b)
            else:
                body.append(b)

        left = []
        right = []
        for b in body:
            x0, x1 = float(b[0]), float(b[2])
            xc = (x0 + x1) / 2.0
            if xc < split_x:
                left.append(b)
            else:
                right.append(b)

        header = _sort_blocks_reading_order(header)
        left = _sort_blocks_reading_order(left)
        right = _sort_blocks_reading_order(right)

        parts = []
        htxt = "\n".join((b[4] or "").strip() for b in header if (b[4] or "").strip()).strip()
        ltxt = "\n".join((b[4] or "").strip() for b in left if (b[4] or "").strip()).strip()
        rtxt = "\n".join((b[4] or "").strip() for b in right if (b[4] or "").strip()).strip()

        for p in (htxt, ltxt, rtxt):
            if p:
                parts.append(p)

        out_pages.append("\n\n".join(parts).strip())

    return "\n\n".join(t for t in out_pages if t.strip()).strip()


def read_pdf_layout_aware(pdf_path: str, min_gap_ratio: float = 0.08) -> str:
    doc = fitz.open(pdf_path)
    try:
        # Decide layout across pages: if ANY page is two-column, treat as two-column.
        two_col_pages = 0
        total_pages = 0
        for page in doc:
            total_pages += 1
            if is_two_column_pymupdf(page, min_gap_ratio=min_gap_ratio):
                two_col_pages += 1

        # if at least 1/3 of pages show two-col, likely a two-col resume
        if total_pages > 0 and (two_col_pages / total_pages) >= 0.34:
            return read_pdf_two_column_left_then_right(doc, min_gap_ratio=min_gap_ratio)

        return read_pdf_sequential(doc)
    finally:
        doc.close()


# ------------------------------------------------------------
# FILE READER (Lambda safe)
# ------------------------------------------------------------

def extract_text_from_file(file_path: str) -> str:
    """
    Supports: .pdf, .docx, .txt

    NOTE: No DOCX->PDF conversion in Lambda using only python-docx + PyMuPDF.
    """
    ext = Path(file_path).suffix.lower()

    try:
        if ext == ".pdf":
            return read_pdf_layout_aware(file_path)

        if ext == ".docx":
            origin = detect_docx_origin(file_path)
            text = read_docx_text(file_path)

            # If it looks like PDF-converted DOCX and text quality is poor, still return docx text
            # (without conversion, this is the best we can do inside Lambda).
            if origin == "PDF_CONVERTED_DOCX" and _looks_bad_text(text):
                logger.warning("DOCX appears PDF-converted and extracted text quality is low; returning best-effort DOCX text.")
            return text

        if ext == ".txt":
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()

        raise ValueError(f"Unsupported file type: {ext}")

    except Exception as e:
        raise Exception(f"Failed to extract text from file: {e}")
